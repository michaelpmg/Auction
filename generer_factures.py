import os
import csv
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from os import listdir
from os.path import isfile, join
from datetime import datetime
from pathvalidate import sanitize_filename

EMAIL_ENCAN_FACTURE = "missladynatalyencan@gmail.com"

def writeTitle(document):
    # Titre "Facture"
    h = document.add_heading("", level=1)
    header_run = h.add_run("FACTURE")
    header_run.bold = True
    paragraph_format = h.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_font = header_run.font
    header_font.size = Pt(24)

def formatFacInfo(info_format, info_run):
    info_format.space_before = Pt(1)
    info_format.space_after = Pt(1)
    info_run.bold = True
    info_run = info_run.font
    info_run.size = Pt(16)
    
def writeNumFac(document, fac_template, num_fac):
    num_fac_p = document.add_paragraph('')
    num_fac_format = num_fac_p.paragraph_format
    num_fac_run = num_fac_p.add_run('Numero de facture : ' + fac_template + str(num_fac))
    formatFacInfo(num_fac_format, num_fac_run)
    num_fac_format.space_before = Pt(10)
    

def writeDate(document, date_str):
    date_p = document.add_paragraph('')
    date_format = date_p.paragraph_format
    date_run = date_p.add_run('date : ' + date_str)
    formatFacInfo(date_format, date_run)

def writeClient(document, client):
    client_p = document.add_paragraph('')
    client_format = client_p.paragraph_format
    client_run = client_p.add_run('Client : ' + client)
    formatFacInfo(client_format, client_run)

def writeDetailHeader(document):
    detail_header_p = document.add_paragraph("")
    detail_header_run = detail_header_p.add_run("Details de la commande")
    detail_header_run.bold = True
    detail_header_font = detail_header_run.font
    detail_header_font.size = Pt(22)
    detail_format = detail_header_p.paragraph_format
    detail_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail_format.space_before = Pt(40)

def writeProductList(document, products):
    fac_total = 0
    
    # list of product
    for product in products:
        prod = products[product]
        prod_qty = prod.qty_
        prod_price = prod.unit_price_
        prod_total = prod_qty * prod_price
        product_str = product + "(x" + str(prod_qty) + " a " + str(prod_price) + "$) = " + str(prod_total) + "$" 

        prod_p = document.add_paragraph('')
        prod_run = prod_p.add_run(product_str)
        prod_format = prod_p.paragraph_format
        formatFacInfo(prod_format, prod_run)
        prod_run.bold = False
        
        fac_total = fac_total + prod_total

    # total
    total_p = document.add_paragraph('')
    total_run = total_p.add_run('Total : ' + str(fac_total) + ".00 $")
    total_run.underline = True
    total_format = total_p.paragraph_format
    formatFacInfo(total_format, total_run)
    total_format.space_before = Pt(10)
    total_font = total_run.font
    total_font.size = Pt(18)
    
def generateDocxForClientSales(client, products, date_str, num_fac, result_file_path):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    fac_id_template = "FAC-" + "".join(date_str.split("-")) + "-"
    
    writeTitle(document)
    writeNumFac(document, fac_id_template, num_fac)
    writeDate(document, date_str)
    writeClient(document, client)
    writeDetailHeader(document)
    writeProductList(document, products)

    # mode paiement
    mode_p = document.add_paragraph('')
    mode_p_format = mode_p.paragraph_format
    mode_p_format.space_before = Pt(40)
    mode_p.paragraph_format.space_after = Pt(2)
    run = mode_p.add_run('Mode de paiement : ')
    run.bold = True
    font = run.font
    font.size = Pt(16)
    run = mode_p.add_run('faire un virement a ' + EMAIL_ENCAN_FACTURE)
    font = run.font
    font.size = Pt(16)

    # password
    pass_p = document.add_paragraph('')
    pass_p.paragraph_format.space_before = Pt(2)
    pass_p.paragraph_format.space_after = Pt(2)
    run = pass_p.add_run('Mot de passe : ')
    run.bold = True
    font = run.font
    font.size = Pt(16)

    run = pass_p.add_run('12345')
    font = run.font
    font.size = Pt(16)

    # status
    status_p = document.add_paragraph('')
    status_p.paragraph_format.space_before = Pt(2)
    run = status_p.add_run('Statut : ')
    run.bold = True
    font = run.font
    font.size = Pt(16)
    
    run = status_p.add_run('Non payé')
    font = run.font
    font.size = Pt(14)

    # skip a line
    document.add_paragraph('')

    # thanks
    thank_p = document.add_paragraph('')
    run = thank_p.add_run('Merci pour votre achat !')
    run.bold = True
    font = run.font
    font.size = Pt(20)
    thank_format = thank_p.paragraph_format
    thank_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.save(result_file_path)
    
def getSalesAsDictFromCSV(file_path):
    sales = []
    with open(file_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            sales.append(row)

    return sales

def ensureDirectoryExists(directory_name):
    os.makedirs(directory_name, exist_ok=True)

class Product:
    def __init__(self):
        self.unit_price_ = 0
        self.qty_ = 0
    
def generateInvoicesFor(cvs_file):
    sales = getSalesAsDictFromCSV(join(to_generate_folder, file))
    
    #get a list of customer from the auction file
    products_by_client = {}
    
    for sale in sales:
        client = sale['Client']
        if client is not "":
            products_by_client[sale['Client']] = {}

    for sale in sales:
        client = sale['Client']
        
        if client == "":
            continue
        
        product_list = {}
        
        #find the client product list or create it
        if products_by_client[client] is not None:
            product_list = products_by_client[client]
        else:
            product_list = {}
            
        #find the current sale product in the product list or create it
        product = Product()
        
        article_name = sale['Article']
        article_qty = 1;
        article_unit_price = 1;

        try:
            article_qty = int(sale['Quantité'])
        except:
            pass

        try:
            article_unit_price = int(sale['Prix Unitaire'])
        except:
            pass
            
        if article_name in product_list:
            product = product_list[article_name]
            product.qty_ = product.qty_ + article_qty
            product.unit_price_ = article_unit_price
        else:
            product = Product()
            product.unit_price_ = article_unit_price
            product.qty_ = article_qty
            
        product_list[article_name] = product
        products_by_client[client] = product_list

    num_client = 0
    print(str(products_by_client))
    for client in products_by_client:
        product_list = products_by_client[client]
        str_now = cvs_file.split('.')[0];
        result_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Factures_' + str_now)
        ensureDirectoryExists(result_dir)
        result_file = os.path.join(result_dir, sanitize_filename(client + "_" + str_now + ".docx"))
        generateDocxForClientSales(client, product_list, str_now, num_client, result_file)
    
def moveEncanFile(csv_file):
    pass
    
if __name__ == "__main__":
    
    to_generate_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Auctions', 'Active')
    files_to_read = [f for f in listdir(to_generate_folder) if isfile(join(to_generate_folder, f)) and f.endswith('.csv')]
    
    for file in files_to_read:
        generateInvoicesFor(file)
